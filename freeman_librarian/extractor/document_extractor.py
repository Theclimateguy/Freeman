"""Read organizational documents and extract canonical entities and relations."""

from __future__ import annotations

from dataclasses import dataclass, field
from hashlib import sha256
from pathlib import Path
import json
import re
from typing import Any, Iterable

from freeman_librarian.extractor.entity_normalizer import EntityNormalizer, infer_kind
from freeman_librarian.ontology.schema import OrganizationDocument, OrganizationEntity, OrganizationRelation, slugify


ENTITY_PREFIX_TO_KIND = {
    "employee": "employee",
    "staff": "employee",
    "role": "role",
    "position": "role",
    "department": "unit",
    "division": "unit",
    "unit": "unit",
    "team": "unit",
    "process": "process",
    "procedure": "process",
    "workflow": "process",
    "document": "document",
    "policy": "document",
    "regulation": "document",
    "artifact": "artifact",
    "system": "system",
    "platform": "system",
}

RELATION_PATTERNS = [
    (re.compile(r"(?P<source>.+?)\s+(owns|is responsible for)\s+(?P<target>.+)", re.IGNORECASE), "owns"),
    (re.compile(r"(?P<source>.+?)\s+participates in\s+(?P<target>.+)", re.IGNORECASE), "participates_in"),
    (re.compile(r"(?P<source>.+?)\s+requires\s+(?P<target>.+)", re.IGNORECASE), "requires"),
    (re.compile(r"(?P<source>.+?)\s+delegates to\s+(?P<target>.+)", re.IGNORECASE), "delegates_to"),
    (re.compile(r"(?P<source>.+?)\s+reports to\s+(?P<target>.+)", re.IGNORECASE), "reports_to"),
]


@dataclass
class ExtractedDocument:
    """Normalized content extracted from one document."""

    document: OrganizationDocument
    entities: list[OrganizationEntity] = field(default_factory=list)
    relations: list[OrganizationRelation] = field(default_factory=list)
    raw_text: str = ""

    def snapshot(self) -> dict[str, Any]:
        return {
            "document": self.document.snapshot(),
            "entities": [entity.snapshot() for entity in self.entities],
            "relations": [relation.snapshot() for relation in self.relations],
            "raw_text": self.raw_text,
        }


class DocumentExtractor:
    """Extract entities and typed relations from a corpus of organizational documents."""

    def __init__(self, *, normalizer: EntityNormalizer | None = None) -> None:
        self.normalizer = normalizer or EntityNormalizer()

    def extract_paths(self, paths: Iterable[str | Path]) -> list[ExtractedDocument]:
        return [self.extract_path(path) for path in paths]

    def extract_path(self, path: str | Path) -> ExtractedDocument:
        source = Path(path).resolve()
        text = self._read_document(source)
        doc_id = f"document_{slugify(source.stem)}"
        document = OrganizationDocument(
            document_id=doc_id,
            title=source.stem,
            path=str(source),
            version="",
            text_hash=sha256(text.encode("utf-8")).hexdigest(),
            metadata={"suffix": source.suffix.lower()},
        )
        entities = self._extract_entities(text, document.document_id)
        relations = self._extract_relations(text, document.document_id, known_entities=entities)
        return ExtractedDocument(document=document, entities=entities, relations=relations, raw_text=text)

    def _read_document(self, path: Path) -> str:
        suffix = path.suffix.lower()
        if suffix in {".txt", ".md", ".rst", ".csv", ".yaml", ".yml"}:
            return path.read_text(encoding="utf-8")
        if suffix == ".json":
            payload = json.loads(path.read_text(encoding="utf-8"))
            return json.dumps(payload, ensure_ascii=False, indent=2, sort_keys=True)
        if suffix == ".pdf":
            try:
                from pypdf import PdfReader
            except ImportError as exc:  # pragma: no cover - optional dependency
                raise RuntimeError("PDF extraction requires `pypdf`. Install with `pip install \"freeman[documents]\"`.") from exc
            return "\n".join(page.extract_text() or "" for page in PdfReader(str(path)).pages)
        if suffix == ".docx":
            try:
                from docx import Document
            except ImportError as exc:  # pragma: no cover - optional dependency
                raise RuntimeError(
                    "DOCX extraction requires `python-docx`. Install with `pip install \"freeman[documents]\"`."
                ) from exc
            return "\n".join(paragraph.text for paragraph in Document(str(path)).paragraphs)
        raise ValueError(f"Unsupported document format: {path.suffix}")

    def _extract_entities(self, text: str, source_id: str) -> list[OrganizationEntity]:
        extracted: dict[str, OrganizationEntity] = {}
        labels_seen: dict[str, str] = {}
        for raw_line in text.splitlines():
            line = raw_line.strip().strip("-*")
            if not line or ":" not in line:
                continue
            prefix, value = [part.strip() for part in line.split(":", 1)]
            kind = ENTITY_PREFIX_TO_KIND.get(prefix.lower())
            if kind is None or not value:
                continue
            entity = self.normalizer.normalize_entity(
                value,
                kind=kind,
                sources=[source_id],
                metadata={"declared_by": prefix.lower()},
            )
            extracted[entity.id] = entity
            labels_seen[entity.label.strip().lower()] = entity.id

        for sentence in self._sentences(text):
            for pattern, relation_type in RELATION_PATTERNS:
                match = pattern.search(sentence)
                if match is None:
                    continue
                for endpoint in ("source", "target"):
                    label = match.group(endpoint).strip(" .;:")
                    inferred_kind = "process" if endpoint == "target" and relation_type in {"owns", "participates_in"} else infer_kind(label)
                    entity = self.normalizer.normalize_entity(
                        label,
                        kind=inferred_kind,
                        sources=[source_id],
                        metadata={"inferred_from_relation": relation_type},
                    )
                    normalized_label = entity.label.strip().lower()
                    if normalized_label in labels_seen:
                        continue
                    extracted.setdefault(entity.id, entity)
                    labels_seen[normalized_label] = entity.id
        return list(extracted.values())

    def _extract_relations(
        self,
        text: str,
        source_id: str,
        *,
        known_entities: Iterable[OrganizationEntity] = (),
    ) -> list[OrganizationRelation]:
        known_by_label = self._known_entity_map(known_entities)
        relations: list[OrganizationRelation] = []
        for sentence in self._sentences(text):
            for pattern, relation_type in RELATION_PATTERNS:
                match = pattern.search(sentence)
                if match is None:
                    continue
                source_label = match.group("source").strip(" .;:")
                target_label = match.group("target").strip(" .;:")
                source_default_kind = "role" if relation_type in {"owns", "participates_in", "delegates_to", "reports_to"} else infer_kind(source_label)
                target_default_kind = "process" if relation_type in {"owns", "participates_in"} else infer_kind(target_label)
                source = self._resolve_entity(source_label, source_id, known_by_label, default_kind=source_default_kind)
                target = self._resolve_entity(target_label, source_id, known_by_label, default_kind=target_default_kind)
                relations.append(
                    OrganizationRelation(
                        source_id=source.id,
                        target_id=target.id,
                        relation_type=relation_type,
                        evidence=[source_id],
                        confidence=0.75,
                        metadata={"sentence": sentence},
                    )
                )
        return relations

    def _known_entity_map(self, entities: Iterable[OrganizationEntity]) -> dict[str, OrganizationEntity]:
        mapping: dict[str, OrganizationEntity] = {}
        for entity in entities:
            mapping[entity.label.strip().lower()] = entity
            for alias in entity.aliases:
                mapping[str(alias).strip().lower()] = entity
        return mapping

    def _resolve_entity(
        self,
        label: str,
        source_id: str,
        known_by_label: dict[str, OrganizationEntity],
        *,
        default_kind: str,
    ) -> OrganizationEntity:
        normalized = label.strip().lower()
        if normalized in known_by_label:
            return known_by_label[normalized]
        return self.normalizer.normalize_entity(label, kind=default_kind, sources=[source_id])

    def _sentences(self, text: str) -> list[str]:
        normalized = re.sub(r"[\r\n]+", ". ", text)
        return [sentence.strip() for sentence in re.split(r"(?<=[.!?])\s+", normalized) if sentence.strip()]


__all__ = ["DocumentExtractor", "ExtractedDocument"]
